import os
try:
    from docx import Document
    print("python-docx is installed.")
except ImportError:
    print("python-docx is NOT installed.")
    exit()

doc_path = "public/Gülnarə m.Mühazirə ERQONOMİKA..docx"

try:
    doc = Document(doc_path)
    print(f"Paragraphs: {len(doc.paragraphs)}")
    text_content = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
    print(f"Text length: {len(text_content)}")
    print("First 500 chars of text:")
    print(text_content[:500])
    
    # Check for images (blip signatures)
    # This is a bit complex in python-docx but we can check relations
    image_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.target_ref:
            image_count += 1
    print(f"Estimated Image count: {image_count}")

except Exception as e:
    print(f"Error reading doc: {e}")
